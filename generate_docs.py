"""
Convert sample .txt contracts into realistic PDF and DOCX files.
Mix of formats to simulate real-world document ingestion.

PDFs:  001, 002, 004, 007, 009 (formal/large contracts tend to be PDF)
DOCXs: 003, 005, 006, 008, 010 (drafts, amendments, smaller contracts)
"""

import os
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from fpdf import FPDF

SRC_DIR = Path("sample_contracts")
OUT_DIR = Path("sample_contracts")

PDF_CONTRACTS = ["001", "002", "004", "007", "009"]
DOCX_CONTRACTS = ["003", "005", "006", "008", "010"]


def sanitize_for_latin1(text: str) -> str:
    """Replace Unicode characters that latin-1 can't encode."""
    replacements = {
        "\u2014": "-",   # em dash
        "\u2013": "-",   # en dash
        "\u2018": "'",   # left single quote
        "\u2019": "'",   # right single quote
        "\u201c": '"',   # left double quote
        "\u201d": '"',   # right double quote
        "\u2026": "...", # ellipsis
        "\u2009": " ",   # thin space
        "\u00e0": "a",   # à (handled by latin-1 but just in case)
        "\u2030": "permille", # per mille
        "\u202f": " ",   # narrow no-break space
        "\u200b": "",    # zero width space
    }
    for char, replacement in replacements.items():
        text = text.replace(char, replacement)
    # Fallback: replace any remaining non-latin-1 chars
    return text.encode("latin-1", errors="replace").decode("latin-1")


# ── PDF Generation ──────────────────────────────────────────────────────────

class ContractPDF(FPDF):
    def __init__(self, title: str):
        super().__init__()
        self._title = title

    def header(self):
        self.set_font("Helvetica", "B", 9)
        self.set_text_color(120, 120, 120)
        self.cell(0, 8, self._title, align="R")
        self.ln(4)
        self.set_draw_color(200, 200, 200)
        self.line(10, self.get_y(), 200, self.get_y())
        self.ln(4)

    def footer(self):
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(150, 150, 150)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")


def make_pdf(txt_path: Path, out_path: Path):
    text = sanitize_for_latin1(txt_path.read_text(encoding="utf-8"))
    lines = text.split("\n")

    # Extract a short title from filename
    name = txt_path.stem.replace("contract_", "").replace("_", " ").title()
    title = sanitize_for_latin1(f"Contract - {name}")

    pdf = ContractPDF(title)
    pdf.alias_nb_pages()
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.add_page()

    for line in lines:
        stripped = line.strip()

        # Section separators
        if stripped.startswith("===="):
            pdf.set_draw_color(0, 51, 102)
            y = pdf.get_y()
            pdf.line(10, y, 200, y)
            pdf.ln(2)
            continue

        # Section headings (all caps lines or ART./ARTICLE/SECTION lines)
        if (stripped.isupper() and len(stripped) > 10) or \
           stripped.startswith("ARTICLE ") or \
           stripped.startswith("ART. ") or \
           stripped.startswith("SECTION ") or \
           stripped.startswith("ARTICOLO ") or \
           stripped.startswith("PREMESSA") or \
           stripped.startswith("BETWEEN:") or \
           stripped.startswith("BETWEEN") or \
           stripped.startswith("TRA:") or \
           stripped.startswith("PARTIES:") or \
           stripped.startswith("FIRME") or \
           stripped.startswith("SIGNATURES"):
            pdf.ln(3)
            pdf.set_font("Helvetica", "B", 11)
            pdf.set_text_color(0, 51, 102)
            pdf.multi_cell(0, 6, stripped)
            pdf.ln(2)
            continue

        # Sub-headings (numbered like 1.1, 2.1, etc.)
        if len(stripped) > 3 and stripped[0].isdigit() and "." in stripped[:4]:
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 5, stripped)
            pdf.ln(1)
            continue

        # Warning/note markers
        if "***" in stripped:
            cleaned = stripped.replace("***", "").strip()
            pdf.set_font("Helvetica", "BI", 9)
            pdf.set_text_color(180, 0, 0)
            pdf.multi_cell(0, 5, cleaned)
            pdf.set_text_color(30, 30, 30)
            pdf.ln(1)
            continue

        # Normal text
        if stripped:
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(30, 30, 30)
            pdf.multi_cell(0, 5, stripped)
            pdf.ln(1)
        else:
            pdf.ln(3)

    pdf.output(str(out_path))
    print(f"  PDF: {out_path}")


# ── DOCX Generation ─────────────────────────────────────────────────────────

def make_docx(txt_path: Path, out_path: Path):
    text = txt_path.read_text(encoding="utf-8")
    lines = text.split("\n")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(30, 30, 30)

    style_heading = doc.styles["Heading 1"]
    style_heading.font.name = "Calibri"
    style_heading.font.size = Pt(14)
    style_heading.font.color.rgb = RGBColor(0, 51, 102)

    style_heading2 = doc.styles["Heading 2"]
    style_heading2.font.name = "Calibri"
    style_heading2.font.size = Pt(11)
    style_heading2.font.color.rgb = RGBColor(0, 51, 102)

    for line in lines:
        stripped = line.strip()

        # Skip separators
        if stripped.startswith("===="):
            continue

        # Main headings
        if (stripped.isupper() and len(stripped) > 10) or \
           stripped.startswith("ARTICLE ") or \
           stripped.startswith("ART. ") or \
           stripped.startswith("SECTION ") or \
           stripped.startswith("ARTICOLO ") or \
           stripped.startswith("PREMESSA") or \
           stripped.startswith("BETWEEN:") or \
           stripped.startswith("TRA:") or \
           stripped.startswith("PARTIES:") or \
           stripped.startswith("FIRME") or \
           stripped.startswith("SIGNATURES"):
            doc.add_heading(stripped, level=1)
            continue

        # Sub-headings
        if len(stripped) > 3 and stripped[0].isdigit() and "." in stripped[:4]:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.bold = True
            run.font.size = Pt(10)
            continue

        # Warnings
        if "***" in stripped:
            cleaned = stripped.replace("***", "").strip()
            p = doc.add_paragraph()
            run = p.add_run(cleaned)
            run.bold = True
            run.italic = True
            run.font.color.rgb = RGBColor(180, 0, 0)
            run.font.size = Pt(9)
            continue

        # Bullet-like lines
        if stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
            continue

        # Normal text or blank
        if stripped:
            doc.add_paragraph(stripped)
        # blank lines are just spacing (Word handles paragraph spacing)

    doc.save(str(out_path))
    print(f"  DOCX: {out_path}")


# ── Main ────────────────────────────────────────────────────────────────────

def main():
    txt_files = sorted(SRC_DIR.glob("contract_*.txt"))

    print(f"Found {len(txt_files)} text contracts. Converting...\n")

    for txt_path in txt_files:
        # Extract contract number (e.g. "001")
        num = txt_path.stem.split("_")[1]

        if num in PDF_CONTRACTS:
            out_path = txt_path.with_suffix(".pdf")
            make_pdf(txt_path, out_path)
        elif num in DOCX_CONTRACTS:
            out_path = txt_path.with_suffix(".docx")
            make_docx(txt_path, out_path)
        else:
            print(f"  SKIP: {txt_path.name} (not in PDF or DOCX list)")

    # Remove .txt originals
    print("\nRemoving .txt originals...")
    for txt_path in txt_files:
        txt_path.unlink()
        print(f"  Removed: {txt_path.name}")

    print("\nDone!")


if __name__ == "__main__":
    main()
